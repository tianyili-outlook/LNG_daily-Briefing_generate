import streamlit as st
import requests
import pandas as pd
import numpy as np
import io
from PyPDF2 import PdfReader
import re
import numbers
from st_aggrid import AgGrid
from datetime import datetime
from datetime import timedelta
from docx import Document
from io import BytesIO

TYPE_TO_SYMBOL = {'JKM': 'AAOVQ00',
                  'HH': 'NMNG001',
                  'NBP': 'AASYR00'}

month_abbr_num = {'Jan': 1,
                  'Feb': 2,
                  'Mar': 3,
                  'Apr': 4,
                  'May': 5,
                  'Jun': 6,
                  'Jul': 7,
                  'Aug': 8,
                  'Sep': 9,
                  'Oct': 10,
                  'Nov': 11,
                  'Dec': 12}

@st.experimental_memo
def get_exchange_rate(date):
    url='http://10.8.19.178:8081/exportData?startTime=&endTime='
    headers={
            "User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/93.0.4577.82 Safari/537.36"
             }
    response = requests.get(url=url,headers=headers, allow_redirects = True).content
    df = pd.read_excel(io.BytesIO(response), index_col = 0)
    return df.loc[date[:4] + '-' + date[4: 6] + '-' + date[6:], '售汇价:美元']

@st.experimental_memo
def get_prices(file, marker_types = TYPE_TO_SYMBOL.keys()):
    price_dict = {marker: [0, 0] for marker in marker_types}
    reader = PdfReader(file)
    page_text = ''
    for i in range(len(reader.pages)):
        page_text += reader.pages[i].extract_text()
    for marker in marker_types:
        price_start_ind = page_text.find(TYPE_TO_SYMBOL[marker])  + len(TYPE_TO_SYMBOL[marker])
        if price_start_ind == -1 + len(TYPE_TO_SYMBOL[marker]):
            marker, ' not found in ', file.name
            'Please manually update price_dict!'
        else:
            temp = page_text[price_start_ind:]
            search = re.search(r'([0-9]*[.])?[0-9]+', temp)
            if search.end() > 10:
                file.name[4:12], marker, ' 价格可能有误，注意手动修改!'
            try:
                price = float(temp[search.start():search.end()])
            except:
                price = 0
                file.name[4:12], marker, ' 价格有误，请手动修改!'
            price_dict[marker][0] = price
            
            ref = page_text.find(TYPE_TO_SYMBOL[marker]) - 12
            temp = page_text[ref:ref+12]
            ref2 = temp.find('(')
            try:
                price_dict[marker][1] = month_abbr_num[temp[ref2 + 1: ref2 + 4]]
            except:
                file.name[4:12], marker, ' 缺失月份数据，请手动修改!'
        result = pd.DataFrame(price_dict, index = ['价格', '月份']).T
        result['种类'] = result.index
        result['月份'] = result['月份'].astype('int32')
        result = result[['种类', '月份', '价格']]
    return result

@st.experimental_memo
def get_pipeline_data(files):
    for file in files:
        if '俄气' in file.name:
            df = pd.read_excel(file)
            i, j = [(i, np.where(df[i] == '本日实际气量')[0].tolist()) for i in list(df) if len(np.where(df[i] == '本日实际气量')[0]) > 0][0]
            i = df.columns.get_loc(i)
            j = j[0]
            mat1 = df.iloc[j+2:j+3, i:i+8].values
        elif '中亚' in file.name:
            df = pd.read_excel(file)
            i, j = [(i, np.where(df[i] == '本日实际气量')[0].tolist()) for i in list(df) if len(np.where(df[i] == '本日实际气量')[0]) > 0][0]
            i = df.columns.get_loc(i)
            j = j[0]
            mat2 = df.iloc[j+2:j+12, i:i+8].values
        elif '缅气' in file.name:
            df = pd.read_excel(file)
            i, j = [(i, np.where(df[i] == '本日实际')[0].tolist()) for i in list(df) if len(np.where(df[i] == '本日实际')[0]) > 0][0]
            i = df.columns.get_loc(i)
            j = j[0]
            mat3 = df.iloc[j+2:j+7, i:i+8].values
        else:
            '文件名中不包含 俄气 中亚 缅气， 请重试'
            return None
    return np.concatenate((mat1, mat2, mat3))

def change_price_describe(document, new_cache_grid):
    def _price_describe_helper(contract_type):
        type_txt_dic = {'HH': ['北美市场：昨日美国假期，纽约商品交易所（NYMEX）HH期货合约无报价。', '北美市场：纽约商品交易所（NYMEX）HH期货合约换至',
                               '北美市场：纽约商品交易所（NYMEX）HH期货合约恢复报价，收于', '北美市场：纽约商品交易所（NYMEX）HH期货合约收于'],
                        'NBP': ['英国市场：昨日英国假期，伦敦洲际交易所（ICE）NBP期货合约无报价。', '英国市场：伦敦洲际交易所（ICE）NBP期货合约换至',
                                '英国市场：伦敦洲际交易所（ICE）NBP期货合约恢复报价，收于', '英国市场：伦敦洲际交易所（ICE）NBP期货合约收于'],
                        'JKM': ['日韩市场：昨日新加坡假期，普氏日韩LNG现货（JKM）无评估。', '日韩市场：普氏日韩LNG现货（JKM）换至',
                                '日韩市场：普氏日韩LNG现货（JKM）恢复评估，价报', '日韩市场：普氏日韩LNG现货（JKM）到岸评估价报']}
        txt = ''
        if new_cache_grid.iloc[-1][contract_type + ' 价格'] == 0:
            txt += type_txt_dic[contract_type][0]
        elif new_cache_grid.iloc[-1][contract_type + ' 月份'] != new_cache_grid.iloc[-2][contract_type + ' 月份']:
            txt += type_txt_dic[contract_type][1] + '{:.0f}月，收于{:.3f}美元/百万英热单位。'.format(new_cache_grid.iloc[-1][contract_type + ' 月份'], 
                                                                                                  new_cache_grid.iloc[-1][contract_type + ' 价格'])
            if contract_type == 'JKM':
                txt = txt.replace('收于', '首日评估')
        elif new_cache_grid.iloc[-2][contract_type + ' 价格'] == 0:
            txt += type_txt_dic[contract_type][2] + '{:.3f}美元/百万英热单位。'.format(new_cache_grid.iloc[-1][contract_type + ' 价格'])
        else:
            price = new_cache_grid.iloc[-1][contract_type + ' 价格']
            diff = price - new_cache_grid.iloc[-2][contract_type + ' 价格']
            change = '上涨' if diff > 0 else '下跌'
            percent = diff / new_cache_grid.iloc[-2][contract_type + ' 价格'] * 100
            if abs(diff) > 1:
                txt += type_txt_dic[contract_type][3] + '{:.3f}美元/百万英热单位，较前一交易日{}{:.3f}美元（{:.2f}%）。'.format(price, change, abs(diff), percent)
            else:
                txt += type_txt_dic[contract_type][3] + '{:.3f}美元/百万英热单位，较前一交易日{}{:.1f}美分（{:.2f}%）。'.format(price, change, abs(diff) * 100, percent)
        return txt
    txt = ''
    for contract_type in ['HH', 'NBP', 'JKM']:
        txt += _price_describe_helper(contract_type)
    document.paragraphs[5].runs[0].text = txt
    return document

def change_news(document, news_title_1, news_content_1, news_title_2, news_content_2):
    document.paragraphs[6].runs[1].text = news_title_1
    document.paragraphs[7].runs[0].text = news_content_1
    document.paragraphs[8].runs[1].text = news_title_2
    document.paragraphs[9].runs[0].text = news_content_2
    return document

def change_misc(document, issue, today, platts_date, maker, checker):
    # change first line issue number, 5 dates, and maker & checker
    # issue
    temp = document.paragraphs[0].runs[3]
    temp.text = issue
    # date 1
    temp = document.paragraphs[2].runs[1]
    temp.text = today[:4]
    temp = document.paragraphs[2].runs[3]
    temp.text = today[4:6].lstrip('0')
    temp = document.paragraphs[2].runs[5]
    temp.text = today[6:].lstrip('0')
    # date 2
    temp = document.paragraphs[4].runs[2]
    temp.text = platts_date[4:6].lstrip('0')
    temp = document.paragraphs[4].runs[4]
    temp.text = platts_date[6:].lstrip('0')
    # date 3
    temp = document.paragraphs[10].runs[1]
    temp.text = platts_date[4:6].lstrip('0')
    temp = document.paragraphs[10].runs[3]
    temp.text = platts_date[6:].lstrip('0')
    # date 4
    temp = document.paragraphs[18].runs[4]
    temp.text = today[:4]
    temp = document.paragraphs[18].runs[6]
    temp.text = today[4:6].lstrip('0')
    temp = document.paragraphs[18].runs[8]
    temp.text = today[6:].lstrip('0')
    # maker
    temp = document.paragraphs[17].runs[1]
    temp.text = maker
    # checker 
    temp = document.paragraphs[17].runs[7]
    temp.text = checker
    return document

def change_table1(document, platts, shangjiao, exchange_rate, today):
    to_fill = np.zeros((7, 3))
    to_fill[0, 0] = platts.iloc[-1]['JKM 价格']
    to_fill[1, 0] = platts.iloc[-1]['HH 价格']
    to_fill[2, 0] = platts.iloc[-2]['HH 价格']
    to_fill[3, 0] = platts['HH 价格'].mean()
    to_fill[4, 0] = platts.iloc[-1]['NBP 价格']
    to_fill[5, 0] = platts.iloc[-2]['NBP 价格']
    to_fill[6, 0] = platts['NBP 价格'].mean()
    to_fill[:, 1] = to_fill[:, 0] * 1000 / 28
    to_fill[:, 2] = to_fill[:, 1] / 1000 * exchange_rate

    temp = document.tables[0].rows[1].cells[0].paragraphs[0].runs[0]
    temp.text = today[:4] + '年' + '{:.0f}'.format(platts.iloc[-1]['JKM 月份']) + '月' + temp.text
    for i in range(1, 8):
        for j in range(2, 5):
            temp = document.tables[0].rows[i].cells[j].paragraphs[0].runs[0]
            temp.text = '{:.2f}'.format(to_fill[i - 1, j - 2])

    to_fill = np.zeros(6)
    to_fill[0] = platts.iloc[-1]['JKM 月份']
    to_fill[1] = platts.iloc[-1]['HH 月份']
    to_fill[2] = platts.iloc[-2]['HH 月份']
    to_fill[4] = platts.iloc[-1]['NBP 月份']
    to_fill[5] = platts.iloc[-2]['NBP 月份']
    for i in [1, 2, 3, 5, 6]:
        temp = document.tables[0].rows[i].cells[5].paragraphs[0].runs[0]
        temp.text = today[:4] + '年' + '{:.0f}'.format(to_fill[i - 1]) + '月' + temp.text

    for i in range(4):
        if shangjiao.iloc[i]['成交量']:
            temp = document.tables[0].rows[i + 8].cells[5].paragraphs[0].runs[0]
            temp.text = shangjiao.iloc[i]['成交量']
            if not shangjiao.iloc[i]['成交价（元/方）']:
                temp = document.tables[0].rows[i + 8].cells[4].paragraphs[0].runs[0]
                temp.text = '成交价未公布'
            else:
                to_fill = np.zeros(3)
                to_fill[2] = float(shangjiao.iloc[i]['成交价（元/方）'])
                to_fill[1] = to_fill[2] / exchange_rate * 1000
                to_fill[0] = to_fill[1] * 28 / 1000
                for j in range(3):
                    temp = document.tables[0].rows[i + 8].cells[j + 2].paragraphs[0].runs[0]
                    temp.text = '{:.2f}'.format(to_fill[j])

    temp = document.paragraphs[11].runs[8]
    temp.text = '{:.4f}'.format(exchange_rate)
    return document

def change_table2(document, pipeline_data, today):
    yesterday = (datetime.strptime(today, '%Y%m%d') - timedelta(1)).strftime('%Y%m%d')
    today_month = today[4:6].lstrip('0')
    today_date = today[6:].lstrip('0')
    yesterday_month = yesterday[4:6].lstrip('0')
    yesterday_date = yesterday[6:].lstrip('0')
    text = '俄气数据截至{t_m}月{t_d}日9:00，中亚气数据截至{y_m}月{y_d}日15:00，缅气数据截至{t_m}月{t_d}日7:30。'.format(
        t_m = today_month,
        t_d = today_date,
        y_m = yesterday_month,
        y_d = yesterday_date)
    temp = document.tables[1].rows[18].cells[0].paragraphs[0]
    temp.runs[0].text = text
    
    for j in [4, 6, 9]:
        temp = document.tables[1].rows[1].cells[j].paragraphs[0].runs[0]
        temp.text = today_month + temp.text

    for i in range(2, 18):
        for j in range(3, 11):
            num = pipeline_data[i - 2, j - 3]
            temp = document.tables[1].rows[i].cells[j].paragraphs[0].runs[0]
            if isinstance(num, numbers.Number) and not np.isnan(num):
                if abs(num) < 1:
                    temp.text = '{:.2f}'.format(num)
                elif abs(num) < 100:
                    temp.text = '{:.1f}'.format(num)
                else:
                    temp.text = '{:,.0f}'.format(num)
            elif isinstance(num, numbers.Number):
                temp.text = ''
            else:
                temp.text = '-'
    return document



with open('cache.txt', 'r') as f:
    issue = int(f.read()) + 1
# issue = 9876

st.title('每日天然气简报生成')
st.header('一、普氏报告处理')
with st.expander('选择1、更新存档'):
    platts_report = st.file_uploader('选择待更新普氏报告：', type = ['pdf'])
    if platts_report is not None:
        date = platts_report.name
        try:
            date = date[date.find('_') + 1: date.find('_') + 9]
            '日期: ', date
        except:
            '普氏报告文件名有误，需保证格式为下划线+8位日期'

        price_df = get_prices(platts_report)
        st.subheader('普氏价格数据： （如有问题请手动在表格中修改）')
        grid_return = AgGrid(price_df, editable=True, height = 200)
        new_price_df = grid_return['data']
    if st.button('刷新存档', key = 1):
        st.session_state.platts_date = date
        df = pd.read_csv('cache.csv', dtype = {'日期': str})
        updated_result = list(new_price_df.values[:, 1:].reshape(-1))
        updated_result.insert(0, date)
        if date in df['日期'].values:
            df[df['日期'] == date] = updated_result
        else:
            df.loc[len(df.index)] = updated_result
        if len(df.index) > 20:
            df = df[-20:]

        new_cache_grid = df.sort_values('日期')
        st.session_state.platts_grid = new_cache_grid
        new_cache_grid.to_csv('cache.csv', index = False)
        '刷新成功！'
        new_cache_grid

with st.expander('选择2、重建存档'):
    platts_reports = st.file_uploader('选择最近20份普氏报告：', type = ['pdf'], accept_multiple_files = True)
    notice = st.text('尚未选取20份报告')
    if len(platts_reports) == 20:
        notice.text('已选择20份报告，开始处理...')
        parse_progress = st.progress(0.0)
        progress_value = 0.0
        values = []
        for platts_report in platts_reports:
            single_result = get_prices(platts_report)
            single_result = list(single_result.values[:, 1:].reshape(-1))
            report_date = platts_report.name
            report_date = report_date[report_date.find('_') + 1: report_date.find('_') + 9]
            single_result.insert(0, report_date)
            values.append(single_result)
            progress_value += 0.05 - 1e-12
            parse_progress.progress(progress_value)
        result = pd.DataFrame(values, columns = ['日期', 'JKM 月份', 'JKM 价格', 'HH 月份', 'HH 价格', 'NBP 月份', 'NBP 价格']).sort_values('日期')
        rebuild_cache_table = AgGrid(result, editable=True, fit_columns_on_grid_load = True)
        new_cache_grid = rebuild_cache_table['data']
        date = new_cache_grid.iloc[-1]['日期']
        
    if st.button('刷新存档', key = 2):
        st.session_state.platts_date = date
        temp = new_cache_grid['日期'].values
        if len(list(temp)) != len(set(temp)):
            '日期有重复，修改后重试'
        else:
            st.session_state.platts_grid = new_cache_grid
            new_cache_grid.to_csv('cache.csv', index = False)
            '刷新成功！'

st.header('二、管道气、汇率等数据获取')
with st.expander('管道气数据'):
    pipeline_files = st.file_uploader('选择中亚、缅气、俄气报表：', type = ['xlsx'], accept_multiple_files = True)
    if len(pipeline_files) == 3:
        st.session_state.pipeline_data = get_pipeline_data(pipeline_files)

with st.expander('汇率'):
    col1, col2, col3 = st.columns([3,2,2], gap = 'large')
    if 'exchange_rate' not in st.session_state:
        st.session_state.exchange_rate = -1
    with col1:
        exchange_date = st.text_input('日期（YYYYMMDD）： ', value = date if 'date' in locals() else '', max_chars = 8)
        if st.button('获取汇率', help = '确保内网链接'):
            st.session_state.exchange_rate = get_exchange_rate(exchange_date)
    with col2:
        alternative_exchange_rate = st.number_input('手动输入汇率： ', min_value = 0., max_value = 70., value = -1., format = '%f')
        if st.button('手动输入汇率', help = '无需内网链接'):
            st.session_state.exchange_rate = alternative_exchange_rate
    with col3:
        '汇率： ', st.session_state.exchange_rate
with st.expander('上交所成交信息'):
    if 'date' in locals():
        url = 'https://weixin.sogou.com/weixin?type=2&s_from=input&query=SHPGX+' + date
    else:
        url = 'https://weixin.sogou.com/weixin?type=2&s_from=input&query=SHPGX'
    st.markdown('[上交所公众号链接](' + url + ')')
    shangjiao_data = pd.DataFrame({'类型': ['PNG挂牌', 'PNG竞价', 'LNG挂牌', 'LNG竞价'], '成交价（元/方）': ['', '', '', ''], '成交量': ['', '', '', '']})
    shangjiao_grid = AgGrid(shangjiao_data, editable=True, fit_columns_on_grid_load = True, height = 200)
    shangjiao_new = shangjiao_grid['data']
    st.session_state.shangjiao = shangjiao_new
with st.expander('新闻翻译'):
    news_title_1 = st.text_input('标题1：')
    news_content_1 = st.text_area('内容1：')
    news_title_2 = st.text_input('标题2：')
    news_content_2 = st.text_area('内容2：')

with st.expander('制表、复核人，期号信息：'):
    col1, col2= st.columns(2, gap = 'large')
    with col1:
        st.text_input('制表人：', key = 'maker')
        new_issue = st.text_input('期号： ', value = str(issue), key = 'new_issue')
    with col2:
        st.text_input('复审人：', key = 'checker')
        st.text_input('日期： ', value = datetime.now().strftime('%Y%m%d'), key = 'today')

st.header('三、简报生成')
document = Document('D:/Python_Script/LNG_daily Briefing_generate/v1.0/template.docx')
target = BytesIO()
col1, col2= st.columns(2, gap = 'large')
with col1:
    if st.button('生成word'):
        with open('cache.txt', 'w') as f:
            f.write(new_issue)

        document = change_price_describe(document, st.session_state.platts_grid)
        document = change_news(document, news_title_1, news_content_1, news_title_2, news_content_2)
        document = change_misc(document, new_issue, st.session_state.today, st.session_state.platts_date, st.session_state.maker, st.session_state.checker)
        document = change_table1(document, st.session_state.platts_grid, st.session_state.shangjiao, st.session_state.exchange_rate, st.session_state.today)
        document = change_table2(document, st.session_state.pipeline_data, st.session_state.today)
        document.save(target)
        '生成完毕！'
with col2:
    st.download_button(label = '下载', 
        data = target, 
        file_name = '每日天然气简报' + st.session_state.today + '.docx')



    